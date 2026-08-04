import json
from html import escape

def export_chat_json(chat, characters, consistency, path):
    payload = {
        "messages": chat,
        "characters": characters,
        "consistency_report": consistency,
    }

    with open(path, "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2, ensure_ascii=False)

def export_chat_docx(chat, characters, consistency, path):
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:
        raise RuntimeError("Install python-docx: pip install python-docx") from exc

    doc = Document()
    doc.add_heading("Aligned Chat Export", 0)

    doc.add_heading("Character Profiles", level=1)
    for char in characters:
        doc.add_heading(char["name"], level=2)
        for key, value in char.items():
            if key != "name":
                p = doc.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))

    doc.add_heading("Consistency Report", level=1)
    for item in consistency:
        doc.add_paragraph(str(item), style=None)

    doc.add_heading("Chat Transcript", level=1)
    for msg in chat:
        p = doc.add_paragraph()
        speaker_run = p.add_run(f'{msg["speaker"]}: ')
        speaker_run.bold = True
        p.add_run(msg["text"])

    for style in doc.styles:
        if style.type == 1:
            style.font.name = "Arial"
            style.font.size = Pt(10)

    doc.save(path)

def export_chat_pdf(chat, characters, consistency, path):
    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
    except Exception as exc:
        raise RuntimeError("Install reportlab: pip install reportlab") from exc

    doc = SimpleDocTemplate(path)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Aligned Chat Export", styles["Title"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Character Profiles", styles["Heading1"]))
    for char in characters:
        story.append(Paragraph(escape(char["name"]), styles["Heading2"]))
        for key, value in char.items():
            if key != "name":
                story.append(Paragraph(f"<b>{escape(str(key))}:</b> {escape(str(value))}", styles["BodyText"]))
        story.append(Spacer(1, 8))

    story.append(Paragraph("Consistency Report", styles["Heading1"]))
    for item in consistency:
        story.append(Paragraph(escape(str(item)), styles["BodyText"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Chat Transcript", styles["Heading1"]))
    for msg in chat:
        line = f"<b>{escape(msg['speaker'])}:</b> {escape(msg['text'])}"
        story.append(Paragraph(line, styles["BodyText"]))
        story.append(Spacer(1, 6))

    doc.build(story)
